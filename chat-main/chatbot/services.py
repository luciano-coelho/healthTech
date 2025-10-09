import smtplib
import time
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

data = {
    "data": "14/09/2024",
    "cliente": "oi",
    "destinatario": "oi",
    "titulo": "Capacitação em Serviços e Produtos para a oi",
    "objetivo": "O principal objetivo dessa capacitação é difundir e repassar aos colaboradores da oi os conhecimentos e habilidades necessárias para aprimorar a prestação de serviços e produtos oferecidos pela empresa.",
    "periodo": "3 meses (previsto)",
    "atribuicao_senac": "atribuicao_senac",
    "atribuicao_cliente": "atribuicao_cliente",
    "detalhamento_proposta": [
        {
            "titulo_etapa": "Módulo 1: Conhecendo os Serviços e Produtos oi",
            "carga_horaria_necessaria": "20 horas",
            "publico_alvo": "Toda a equipe de atendimento e vendas",
            "objetivo_da_etapa": "Apresentar e detalhar os serviços e produtos oferecidos pela oi, bem como as suas características, funcionalidades e benefícios.",
            "conteudo_programatico": [
                "Histórico e evolução da oi",
                "Portfólio completo de serviços e produtos",
                "Análise do mercado e da concorrência",
                "Estratégias de vendas e atendimento"
            ]
        },
        {
            "titulo_etapa": "Módulo 2: Técnicas de Atendimento e Vendas",
            "carga_horaria_necessaria": "40 horas",
            "publico_alvo": "Equipe de atendimento e vendas",
            "objetivo_da_etapa": "Desenvolver habilidades de atendimento ao cliente e técnicas de vendas eficazes, com foco na satisfação do cliente.",
            "conteudo_programatico": [
                "Técnicas de comunicação e relacionamento interpessoal",
                "Gestão de conflitos e reclamações",
                "Identificação e antecipação das necessidades do cliente",
                "Técnicas de negociação e fechamento de vendas"
            ]
        },
        {
            "titulo_etapa": "Módulo 3: Gestão de Serviços e Produtos",
            "carga_horaria_necessaria": "20 horas",
            "publico_alvo": "Equipes de gestão e supervisão",
            "objetivo_da_etapa": "Capacitar gestores e supervisores na gestão eficaz de serviços e produtos, otimizando processos e garantindo a qualidade.",
            "conteudo_programatico": [
                "Planejamento e organização de serviços e produtos",
                "Gestão de equipes e desempenho",
                "Controle de qualidade e indicadores de desempenho",
                "Atendimento a reclamações e resolução de problemas"
            ]
        }
    ],
    "carga_horaria_total": "80 horas",
    "valor_investimento": "R$ 10.000,00",
    "type": "public"
}

REPLACEMENTS = {
    "{data}": "data",
    "{cliente}": "cliente",
    "{destinatario}": "destinatario",
    "{titulo}": "titulo",
    "{objetivo}": "objetivo",
    "{periodo}": "periodo",
    "{carga_horaria_total}": "carga_horaria_total",
    "{valor_investimento}": "valor_investimento",
    "{type}": "type",
    "{atribuicao_senac}": "atribuicao_senac",
    "{atribuicao_cliente}": "atribuicao_cliente",
}

SCHEMA_DETALHAMENTO_PROPOSTA = {
    "carga_horaria_necessaria": "Carga Horária",
    "publico_alvo": "Público-alvo",
    "objetivo_da_etapa": "Objetivo",
    "conteudo_programatico": "Conteúdo Programático (ementa)",
}


def replace_text_in_doc(doc, data_json):
    replace_paragraph(doc, data_json)
    if data_json.get("detalhamento_proposta"):
        insert_detalhamento_da_proposta(
            doc, data_json.get("detalhamento_proposta"))


def replace_paragraph(doc, data_json):
    for paragraph in doc.paragraphs:
        for key, value in REPLACEMENTS.items():
            if paragraph.text == '{type}':
                insert_private_or_public(paragraph, data_json["type"])
                break
            if key in paragraph.text:
                text_replace = data_json[value] if data_json[value] else ''
                paragraph.text = paragraph.text.replace(key, text_replace)


def insert_detalhamento_da_proposta(doc, detalhamento_proposta: list):
    table = doc.tables[2]
    for detalhe in detalhamento_proposta:
        for key, value in detalhe.items():
            if key == "titulo_etapa":
                insert_line_merge(table, value)
                continue
            insert_line_information(
                table, SCHEMA_DETALHAMENTO_PROPOSTA[key], value)

    first_row = table.rows[0]
    merged_cell = first_row.cells[0].merge(first_row.cells[1])
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                            bottom={"sz": 12, "val": "single",
                                    "color": "000000"},
                            left={"sz": 12, "val": "single", "color": "000000"},
                            right={"sz": 12, "val": "single", "color": "000000"})


def insert_line_merge(table, text):
    merged_row = table.add_row()
    merged_row.cells[0].text = text
    merged_cell = merged_row.cells[0].merge(merged_row.cells[1])
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].runs[0].bold = True


def insert_line_information(table, key, value):
    new_row = table.add_row()

    new_row.cells[0].text = key
    if isinstance(value, list):
        for item in value:
            new_row.cells[1].add_paragraph(f"• {item}")
    else:
        new_row.cells[1].text = value

    table_width = Inches(6)
    cell_1_width = int(table_width * 0.2)
    cell_2_width = int(table_width * 0.8)
    new_row.cells[0].width = cell_1_width
    new_row.cells[1].width = cell_2_width


def set_cell_border(cell, **kwargs):
    """
    Define bordas para uma célula. Os argumentos aceitos em kwargs são:
    top, left, bottom, right e os valores podem ser:
    {"sz": tamanho, "val": "single", "color": "000000", "space": "0"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Criar o elemento <w:tcBorders> se não existir
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Definir as bordas para cada lado (topo, esquerdo, inferior, direito)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get("val", "single"))
            # Tamanho da borda (4 = 1/2 pt)
            element.set(qn('w:sz'), str(kwargs[edge].get("sz", 4)))
            element.set(qn('w:space'), str(kwargs[edge].get("space", 0)))
            element.set(qn('w:color'), kwargs[edge].get(
                "color", "000000"))  # Cor da borda
            tcBorders.append(element)


def insert_private_or_public(paragraph, type: str):
    if type == "private":
        texts_insert = [
            ("Empresa privada: ", True),  # Negrito
            ("      • Ato Constitutivo, Estatuto ou Contrato Social em vigor, acompanhado da última Alteração Contratual, ou a última Alteração Contratual Consolidada, se houver devidamente registrados.\n", False),
            ("      • Ata de eleição da Diretoria e/ou Conselho de Administração, quando se aplicar.\n", False),
            ("      • Procuração, quando necessário.\n", False),
            ("      • Cópia do comprovante de inscrição no cadastro nacional de pessoas jurídicas (CNPJ).\n", False),
            ("      • Cópia da Identidade e do CPF do representante legal / administrador.\n", False),
            ("      • Nome, CPF e e-mail individual do representante legal e da testemunha que assinará o instrumento contratual.\n", False),
            ("      • Proposta final apresentada e aprovada.\n", False),
            ("      • Data dos pagamentos de acordo com os possíveis parcelamentos.\n", False)
        ]
    else:

        texts_insert = [
            ("Ente público:\n", True),  # Negrito
            ("      • Minuta do Contrato emitida pelo Ente Público\n", False),
            ("      • Ato Administrativo de nomeação da autoridade máxima do Ente Público.\n", False),
            ("      • Ato Administrativo para designação de responsável autorizando a delegação de atribuições para assinatura de contratos/convênios.\n", False),
            ("      • Nome, CPF e e-mail individual do representante legal e da testemunha que assinará o instrumento contratual, caso o ente público aceite.\n", False),
            ("      • Proposta final apresentada e aprovada.\n", False),
            ("      • Data dos pagamentos de acordo com os possíveis parcelamentos.\n", False)
        ]

    paragraph.clear()
    for text, is_bold in texts_insert:
        run = paragraph.add_run(text)
        run.bold = is_bold
        if is_bold:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def generate_doc(data_json):
    doc = Document("modelo.docx")
    replace_text_in_doc(doc, data_json)
    file_name = f"{time.time()}.docx"
    doc.save(file_name)
    content_file = open(file_name, "rb")
    return content_file


def send_email(to_email, subject, file_path=None,  body=""):

    # onfigs
    from_email = "geniapropostas@gmail.com"
    password = "hkjb rvli mzfq eaqb"

    # Configurações do servidor SMTP (exemplo com Gmail)
    smtp_server = "smtp.gmail.com"
    smtp_port = 587  # Porta TLS

    # Criação da mensagem
    msg = MIMEMultipart()
    msg["From"] = from_email
    msg["To"] = to_email
    msg["Subject"] = subject

    # Corpo do e-mail
    msg.attach(MIMEText(body, "plain"))

    # Se houver conteúdo de arquivo e nome do arquivo
    if file_path:

        try:
            with open(file_path, "rb") as f:
                file_content = f.read()

            file_name = file_path
            # Criar o objeto MIMEBase e adicionar o cabeçalho adequado
            part = MIMEBase("application", "octet-stream")
            part.set_payload(file_content)
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {file_name}",
            )

            # Anexar o arquivo à mensagem
            msg.attach(part)
        except Exception as e:
            print(f"Erro ao anexar o arquivo: {e}")
            return

    try:
        # Conectar ao servidor SMTP
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()  # Inicializa a conexão TLS
        server.login(from_email, password)  # Autentica no servidor SMTP

        # Enviar o e-mail
        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        print(f"E-mail enviado com sucesso para {to_email}")

    except Exception as e:
        print(f"Erro ao enviar o e-mail: {e}")

    finally:
        # Fechar a conexão SMTP
        server.quit()


if __name__ == '__main__':
    with open("modelo.docx", "rb") as f:
        file_content = f.read()

    file_name = "arquivo.txt"

    send_email_with_attachment(
        to_email="alexsandervieirajunior@gmail.com",
        subject="Proposta Gerada GenIA",
        file_content=file_content,
        file_name=file_name
    )
